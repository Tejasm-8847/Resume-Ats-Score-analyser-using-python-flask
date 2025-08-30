from flask import Flask, render_template, request, redirect, url_for, flash, session, send_file
from flask_sqlalchemy import SQLAlchemy
from flask_login import LoginManager, UserMixin, login_user, logout_user, login_required, current_user
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
import os
import tempfile
from datetime import datetime
import uuid
import re
from io import BytesIO

# Document processing libraries
import docx
import PyPDF2
from docx import Document

app = Flask(__name__)
app.config['SECRET_KEY'] = 'your_secret_key'
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///resume_optimizer.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max upload

# Create uploads folder if it doesn't exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

db = SQLAlchemy(app)
login_manager = LoginManager(app)
login_manager.login_view = 'login'

ALLOWED_EXTENSIONS = {'pdf', 'docx'}

# Models
class User(UserMixin, db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(100), unique=True, nullable=False)
    email = db.Column(db.String(100), unique=True, nullable=False)
    password_hash = db.Column(db.String(200), nullable=False)
    resumes = db.relationship('Resume', backref='user', lazy=True)
    
    def set_password(self, password):
        self.password_hash = generate_password_hash(password)
        
    def check_password(self, password):
        return check_password_hash(self.password_hash, password)
    
class Resume(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    filename = db.Column(db.String(200), nullable=False)
    original_text = db.Column(db.Text, nullable=True)
    optimized_text = db.Column(db.Text, nullable=True)
    job_description = db.Column(db.Text, nullable=True)
    ats_score = db.Column(db.Float, nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    file_type = db.Column(db.String(10), nullable=False)
    
@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

# Helper functions
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_pdf(pdf_file):
    pdf_reader = PyPDF2.PdfReader(pdf_file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text() if page.extract_text() else ""
    return text

def extract_text_from_docx(docx_file):
    doc = Document(docx_file)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text

def parse_resume(file_path, file_type):
    """Extract text from resume file"""
    if file_type == 'pdf':
        with open(file_path, 'rb') as file:
            return extract_text_from_pdf(file)
    elif file_type == 'docx':
        return extract_text_from_docx(file_path)
    return ""

def compare_with_job_description(resume_text, job_description):
    """Compare resume with job description to find missing keywords"""
    resume_words = set(re.findall(r'\b\w+\b', resume_text.lower()))
    job_words = set(re.findall(r'\b\w+\b', job_description.lower()))
    
    # Simple keyword matching (would be more sophisticated in production)
    important_words = job_words - resume_words
    
    # Calculate basic ATS score (percentage of job keywords found in resume)
    if len(job_words) > 0:
        match_words = job_words.intersection(resume_words)
        score = (len(match_words) / len(job_words)) * 100
    else:
        score = 0
        
    return {
        'score': score,
        'missing_keywords': list(important_words)
    }

def optimize_resume(resume_text, job_description, analysis):
    """Generate an optimized version of the resume based on the job description"""
    # In a real system, this would be more sophisticated
    # For demo purposes, we'll just add a section with missing keywords
    
    optimized = resume_text
    
    if analysis['missing_keywords']:
        missing_keywords = ", ".join(analysis['missing_keywords'][:10])  # First 10 keywords
        suggestion = f"\n\n--- OPTIMIZATION SUGGESTIONS ---\n"
        suggestion += f"Consider adding these keywords: {missing_keywords}\n"
        suggestion += f"Your resume matches {analysis['score']:.1f}% of job requirements.\n"
        
        optimized += suggestion
        
    return optimized

def create_optimized_docx(original_text, optimized_text):
    """Create a new DOCX with the optimized content"""
    doc = Document()
    doc.add_heading('Optimized Resume', 0)
    
    # Add the optimized text
    paragraphs = optimized_text.split('\n')
    for para in paragraphs:
        if para.strip():
            doc.add_paragraph(para)
    
    # Save to a BytesIO object
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Routes
@app.route('/')
def index():
    return render_template('index.html')

@app.route('/register', methods=['GET', 'POST'])
def register():
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
    
    if request.method == 'POST':
        username = request.form.get('username')
        email = request.form.get('email')
        password = request.form.get('password')
        
        # Check if username or email already exists
        if User.query.filter_by(username=username).first():
            flash('Username already exists.', 'danger')
            return redirect(url_for('register'))
        
        if User.query.filter_by(email=email).first():
            flash('Email already registered.', 'danger')
            return redirect(url_for('register'))
        
        # Create new user
        new_user = User(username=username, email=email)
        new_user.set_password(password)
        
        db.session.add(new_user)
        db.session.commit()
        
        flash('Registration successful! Please log in.', 'success')
        return redirect(url_for('login'))
    
    return render_template('register.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
    
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        
        user = User.query.filter_by(username=username).first()
        
        if user and user.check_password(password):
            login_user(user)
            next_page = request.args.get('next')
            return redirect(next_page or url_for('dashboard'))
        else:
            flash('Invalid username or password', 'danger')
            
    return render_template('login.html')

@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('index'))

@app.route('/dashboard')
@login_required
def dashboard():
    user_resumes = Resume.query.filter_by(user_id=current_user.id).all()
    return render_template('dashboard.html', resumes=user_resumes)

@app.route('/upload', methods=['GET', 'POST'])
@login_required
def upload_resume():
    if request.method == 'POST':
        # Check if resume file was uploaded
        if 'resume' not in request.files:
            flash('No file selected', 'danger')
            return redirect(request.url)
        
        file = request.files['resume']
        job_description = request.form.get('job_description', '')
        
        if file.filename == '':
            flash('No file selected', 'danger')
            return redirect(request.url)
        
        if file and allowed_file(file.filename):
            # Save file with secure filename
            filename = secure_filename(file.filename)
            file_extension = filename.rsplit('.', 1)[1].lower()
            unique_filename = f"{uuid.uuid4().hex}.{file_extension}"
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)
            file.save(file_path)
            
            # Parse resume
            resume_text = parse_resume(file_path, file_extension)
            
            # Only analyze if job description is provided
            if job_description:
                analysis = compare_with_job_description(resume_text, job_description)
                optimized_text = optimize_resume(resume_text, job_description, analysis)
                ats_score = analysis['score']
            else:
                optimized_text = resume_text
                ats_score = 0
            
            # Save to database
            new_resume = Resume(
                filename=filename,
                original_text=resume_text,
                optimized_text=optimized_text,
                job_description=job_description,
                ats_score=ats_score,
                user_id=current_user.id,
                file_type=file_extension
            )
            
            db.session.add(new_resume)
            db.session.commit()
            
            flash('Resume uploaded and analyzed successfully!', 'success')
            return redirect(url_for('view_resume', resume_id=new_resume.id))
        else:
            flash('Only PDF and DOCX files are allowed', 'danger')
            
    return render_template('upload.html')

@app.route('/resume/<int:resume_id>')
@login_required
def view_resume(resume_id):
    resume = Resume.query.get_or_404(resume_id)
    
    # Make sure the resume belongs to the current user
    if resume.user_id != current_user.id:
        flash('Access denied', 'danger')
        return redirect(url_for('dashboard'))
    
    return render_template('view_resume.html', resume=resume)

@app.route('/download/<int:resume_id>')
@login_required
def download_resume(resume_id):
    resume = Resume.query.get_or_404(resume_id)
    
    # Make sure the resume belongs to the current user
    if resume.user_id != current_user.id:
        flash('Access denied', 'danger')
        return redirect(url_for('dashboard'))
    
    # Create optimized docx file
    buffer = create_optimized_docx(resume.original_text, resume.optimized_text)
    
    # Return the file
    return send_file(
        buffer,
        as_attachment=True,
        download_name=f"optimized_{resume.filename.rsplit('.', 1)[0]}.docx",
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

# Initialize database
with app.app_context():
    db.create_all()

if __name__ == "__main__":
    app.run(debug=True)

